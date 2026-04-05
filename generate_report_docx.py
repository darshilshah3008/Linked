"""Generate a comprehensive dashboard report as a Word document from dashboard.json.

Covers ALL companies, ALL leads, ALL contacts -- grouped and detailed.
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def _set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, size=9)
        shading = table.rows[0].cells[i]._element
        from docx.oxml.ns import qn
        from lxml import etree
        tc_pr = shading.get_or_add_tcPr()
        shd = etree.SubElement(tc_pr, qn("w:shd"))
        shd.set(qn("w:fill"), "2E4057")
        shd.set(qn("w:val"), "clear")
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=9)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def _relevance_label(rel):
    return (rel or "unknown").upper()


def _step_label(step):
    return (step or "monitor_weekly").replace("_", " ").title()


def main():
    with open("dashboard.json", "r", encoding="utf-8") as f:
        d = json.load(f)

    doc = Document()

    # ── Title ─────────────────────────────────────────────
    title = doc.add_heading("Comprehensive Job Search Dashboard", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {d['date']}")
    doc.add_paragraph(f"Summary: {d['summary']}")

    # ── Quick stats ───────────────────────────────────────
    n_jobs = len(d["top_jobs"])
    n_people = len(d["top_people"])
    n_companies = len(d["top_companies"])
    n_outreach = len(d["outreach_drafts"])

    # Group companies by relevance for stats
    co_by_rel = {}
    for co in d["top_companies"]:
        rel = (co.get("embedded_relevance") or "unknown").lower()
        co_by_rel.setdefault(rel, []).append(co)

    stats_para = doc.add_paragraph()
    stats_para.add_run(f"Total Job Leads: {n_jobs}   |   ").bold = True
    stats_para.add_run(f"Contacts: {n_people}   |   ").bold = True
    stats_para.add_run(f"Companies: {n_companies}   |   ").bold = True
    stats_para.add_run(f"Outreach Drafts: {n_outreach}").bold = True
    doc.add_paragraph(
        f"Companies by Relevance:  HIGH: {len(co_by_rel.get('high', []))}  |  "
        f"MEDIUM: {len(co_by_rel.get('medium', []))}  |  "
        f"LOW: {len(co_by_rel.get('low', []))}  |  "
        f"UNKNOWN: {len(co_by_rel.get('unknown', []))}"
    )
    doc.add_paragraph("")

    # ══════════════════════════════════════════════════════
    # SECTION 1: ALL JOB LEADS (sorted by fit score)
    # ══════════════════════════════════════════════════════
    doc.add_heading(f"Section 1: All {n_jobs} Job Leads (Ranked by Fit Score)", level=1)

    headers = ["#", "Title", "Company", "Location", "Fit", "Confidence"]
    rows = []
    for i, job in enumerate(d["top_jobs"], 1):
        rows.append([
            str(i),
            job["title"],
            job["company"],
            job["location"],
            f"{job['fit_score']:.0f}/100",
            job["confidence"],
        ])
    _add_table(doc, headers, rows, col_widths=[0.3, 2.0, 1.3, 1.0, 0.5, 0.7])

    # Detailed breakdown per job
    doc.add_paragraph("")
    doc.add_heading("Detailed Job Breakdown", level=2)

    for i, job in enumerate(d["top_jobs"], 1):
        doc.add_heading(f"{i}. {job['title']} at {job['company']}", level=3)
        doc.add_paragraph(f"Location: {job['location']}  |  Fit Score: {job['fit_score']:.0f}/100  |  Confidence: {job['confidence']}")
        doc.add_paragraph(f"Apply URL: {job['url']}")
        doc.add_paragraph(f"Description: {job['description']}")

        if job.get("match_reasons"):
            p = doc.add_paragraph("Match Reasons:", style="List Bullet")
            p.runs[0].bold = True
            for reason in job["match_reasons"]:
                doc.add_paragraph(reason, style="List Bullet 2")

        missing = job.get("missing_requirements", [])
        if missing:
            p = doc.add_paragraph("Gaps:", style="List Bullet")
            p.runs[0].bold = True
            for m in missing:
                doc.add_paragraph(m, style="List Bullet 2")

    # ══════════════════════════════════════════════════════
    # SECTION 2: ALL CONTACTS
    # ══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading(f"Section 2: All {n_people} Contacts", level=1)

    headers = ["#", "Name", "Role", "Company", "Priority", "Outreach Type"]
    rows = []
    for i, p in enumerate(d["top_people"], 1):
        outreach = (p.get("suggested_outreach_type") or "connection").replace("_", " ").title()
        rows.append([
            str(i),
            p["name"],
            p["role"],
            p["company"],
            (p.get("contact_priority") or "medium").upper(),
            outreach,
        ])
    _add_table(doc, headers, rows, col_widths=[0.3, 1.5, 1.5, 1.2, 0.7, 1.0])

    doc.add_paragraph("")
    for i, p in enumerate(d["top_people"], 1):
        doc.add_heading(f"{i}. {p['name']}", level=3)
        outreach = (p.get("suggested_outreach_type") or "connection").replace("_", " ").title()
        priority = (p.get("contact_priority") or "medium").upper()
        doc.add_paragraph(f"Role: {p['role']}  |  Company: {p['company']}")
        doc.add_paragraph(f"Priority: {priority}  |  Outreach: {outreach}")
        doc.add_paragraph(f"LinkedIn: {p.get('linkedin_url') or 'N/A'}")
        doc.add_paragraph(f"Relevance: {p.get('relevance_reason', 'N/A')}")

    # ══════════════════════════════════════════════════════
    # SECTION 3: ALL COMPANIES (grouped by relevance)
    # ══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading(f"Section 3: All {n_companies} Target Companies", level=1)

    # Master table with all companies
    headers = ["#", "Company", "Industry", "Relevance", "Next Step", "Careers URL"]
    rows = []
    for i, co in enumerate(d["top_companies"], 1):
        rows.append([
            str(i),
            co["name"],
            co.get("industry", "N/A"),
            _relevance_label(co.get("embedded_relevance")),
            _step_label(co.get("suggested_next_step")),
            co.get("careers_url", "N/A"),
        ])
    _add_table(doc, headers, rows, col_widths=[0.3, 1.3, 1.4, 0.7, 0.9, 2.0])

    # Detailed breakdown grouped by relevance
    for rel_level in ["high", "medium", "low", "unknown"]:
        group = co_by_rel.get(rel_level, [])
        if not group:
            continue
        doc.add_page_break()
        doc.add_heading(f"Companies - {rel_level.upper()} Relevance ({len(group)} companies)", level=2)

        for co in group:
            doc.add_heading(f"{co['name']}", level=3)
            doc.add_paragraph(f"Industry: {co.get('industry', 'N/A')}")
            doc.add_paragraph(f"Embedded Relevance: {_relevance_label(co.get('embedded_relevance'))}")
            doc.add_paragraph(f"Careers URL: {co.get('careers_url', 'N/A')}")
            doc.add_paragraph(f"Summary: {co.get('summary', 'N/A')}")
            doc.add_paragraph(f"Research Notes: {co.get('research_notes', 'N/A')}")
            doc.add_paragraph(f"Recommended Action: {_step_label(co.get('suggested_next_step'))}")

    # ══════════════════════════════════════════════════════
    # SECTION 4: OUTREACH DRAFTS
    # ══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading(f"Section 4: Outreach Drafts ({n_outreach} Total)", level=1)

    for i, msg in enumerate(d["outreach_drafts"], 1):
        mtype = msg["message_type"].replace("_", " ").title()
        doc.add_heading(f"{i}. {mtype} - {msg['company']}", level=3)
        doc.add_paragraph(f"Target Role: {msg['lead_title']}")
        doc.add_paragraph(f"Message Type: {mtype}")

        draft_para = doc.add_paragraph()
        draft_para.paragraph_format.left_indent = Inches(0.5)
        for line in msg["draft_text"].split("\n"):
            if draft_para.text:
                draft_para.add_run("\n")
            run = draft_para.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(44, 62, 80)

    # ══════════════════════════════════════════════════════
    # SECTION 5: LINKEDIN COMMENTS
    # ══════════════════════════════════════════════════════
    doc.add_page_break()
    n_comments = len(d["comments_to_write"])
    doc.add_heading(f"Section 5: LinkedIn Comments ({n_comments} Total)", level=1)

    for i, c in enumerate(d["comments_to_write"], 1):
        doc.add_heading(f"{i}. Theme: {c['theme']}", level=3)
        p = doc.add_paragraph()
        run = p.add_run(f"Hook: {c['hook']}")
        run.bold = True
        run.font.size = Pt(10)

        if c.get("full_text"):
            doc.add_paragraph("")
            draft_para = doc.add_paragraph()
            draft_para.paragraph_format.left_indent = Inches(0.5)
            for line in c["full_text"].split("\n"):
                if draft_para.text:
                    draft_para.add_run("\n")
                run = draft_para.add_run(line)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(44, 62, 80)

    # ══════════════════════════════════════════════════════
    # SECTION 6: LINKEDIN POST
    # ══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Section 6: LinkedIn Post Draft", level=1)

    lp = d.get("linkedin_post_draft")
    if lp:
        doc.add_paragraph(f"Theme: {lp['theme']}")
        p = doc.add_paragraph()
        run = p.add_run(f"Hook: {lp['hook']}")
        run.bold = True

        if lp.get("outline"):
            doc.add_paragraph(f"Outline: {lp['outline']}")

        if lp.get("full_text"):
            doc.add_heading("Full Post", level=3)
            draft_para = doc.add_paragraph()
            draft_para.paragraph_format.left_indent = Inches(0.5)
            for line in lp["full_text"].split("\n"):
                if draft_para.text:
                    draft_para.add_run("\n")
                run = draft_para.add_run(line)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(44, 62, 80)

    # ── Footer ────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"End of Report  |  {n_companies} companies  |  {n_jobs} leads  |  {n_people} contacts"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    out_path = Path("dashboard_report.docx")
    try:
        doc.save(str(out_path))
    except PermissionError:
        from datetime import datetime
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_path = Path(f"dashboard_report_{ts}.docx")
        doc.save(str(out_path))
    print(f"Report saved: {out_path} ({out_path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
